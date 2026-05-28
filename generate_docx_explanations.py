from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = r'C:\Users\Lenovo\Documents\AI in Bioinformatics\Project_Explanations'

def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    return h

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_step(doc, step_num, title, explanation, function_name, why_text):
    add_h(doc, f"Step {step_num}: {title}", level=2)
    
    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run(explanation)
    
    p2 = doc.add_paragraph()
    run2 = p2.add_run("Why we do this: ")
    run2.bold = True
    run2.font.color.rgb = RGBColor(0xE7, 0x45, 0x60)
    p2.add_run(why_text)
    
    if function_name:
        p3 = doc.add_paragraph()
        run3 = p3.add_run("Function(s) used: ")
        run3.bold = True
        p3.add_run(function_name)

def add_code(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    return p


# ════════════════════════════════════════════════════════════════════
# PROJECT 1
# ════════════════════════════════════════════════════════════════════
doc1 = Document()
doc1.add_heading('Project 1: DNA Sequence Classification', 0)
add_para(doc1, 'A step-by-step beginner-friendly walkthrough of how we classify DNA sequences using k-mer features and Random Forest.')

add_h(doc1, "The Big Picture — What & Why")
add_para(doc1, "DNA sequences from different species look different. Some species have lots of G and C (high GC), others have more A and T (low GC). We want to teach a computer to look at a piece of DNA and tell us which species it came from.")
add_para(doc1, "Why? In real life, scientists find DNA fragments in soil, water, or patient samples and need to know what organism they belong to — this is called metagenomic binning. It's how we discover new viruses, track outbreaks, and study microbiomes.")

add_step(doc1, 1, "Generate Synthetic DNA Sequences",
    "We create fake DNA sequences with controlled GC content. Species A gets 70% GC (like Streptomyces bacteria), Species B gets 25% GC (like Plasmodium parasite). Each sequence is 500 letters (A, T, G, C) long.",
    "generate_dna_sequence(length, gc_content)",
    "We start with synthetic data so we KNOW the correct answer. This lets us check if our method works before trying it on real, messy data. Real data from NCBI GenBank can replace this later.")

add_step(doc1, 2, "Extract k-mer Features",
    "We chop each DNA sequence into overlapping 4-letter chunks (k=4). For example, 'ATGCGTA' becomes ATGC, TGCG, GCGT, CGTA. We count how many times each possible 4-letter combination appears, giving 4^4 = 256 numbers per sequence.",
    "extract_kmer_features(sequence, k=4)",
    "Computers can't read DNA letters directly — they need numbers. k-mer counting converts a variable-length DNA string into a fixed-size numerical vector. This is the most common way to featurize DNA for ML.")

add_step(doc1, 3, "Split Data into Train/Test",
    "We randomly split our data: 75% goes to training (the computer learns from this), 25% goes to testing (we check if it learned correctly). The split keeps the same ratio of Species A and B in both sets.",
    "train_test_split(X, y, test_size=0.25, stratify=y)",
    "If we test on the same data the model trained on, we'd cheat — it would just memorize answers. Testing on held-out data tells us if the model can generalize to NEW sequences it has never seen.")

add_step(doc1, 4, "Train Random Forest Classifier",
    "Random Forest builds 200 different decision trees, each trained on a random subset of features and samples. Each tree makes a vote, and the majority decides the species. Think of it like asking 200 experts instead of just one.",
    "RandomForestClassifier(n_estimators=200, max_depth=20)",
    "Random Forest is great for beginners because: (1) it works well with many features, (2) it's hard to overfit, (3) it tells us which features matter most, (4) it handles the sparse k-mer data well.")

add_step(doc1, 5, "Evaluate the Model",
    "We compare the model's predictions against the true labels on the test set. We check: Accuracy (what fraction correct?), ROC-AUC (how well does it separate classes?), and the Confusion Matrix (where does it get confused?).",
    "accuracy_score(), roc_auc_score(), confusion_matrix(), classification_report()",
    "Accuracy alone can be misleading (e.g., 95% accuracy on imbalanced data). ROC-AUC tells us about separability. The confusion matrix shows if Species A is ever mistaken for B. Feature importance reveals which k-mers drive the decision.")

add_h(doc1, "What To Expect")
add_para(doc1, "With 70% vs 25% GC, the model should achieve ~99% accuracy. The top features will be GC-rich k-mers (GGCC, CCGG) for Species A and AT-rich k-mers (AATT, TTAA) for Species B. The output includes a confusion matrix plot and ROC curve saved as project1_confusion_roc.png.")

doc1.save(os.path.join(OUTPUT_DIR, 'Project1_DNA_kmer_Explained.docx'))
print("1/4 done")


# ════════════════════════════════════════════════════════════════════
# PROJECT 2
# ════════════════════════════════════════════════════════════════════
doc2 = Document()
doc2.add_heading('Project 2: Protein Family Classification with 1D CNN', 0)
add_para(doc2, 'A step-by-step walkthrough of how we classify proteins into families using a Convolutional Neural Network.')

add_h(doc2, "The Big Picture — What & Why")
add_para(doc2, "Proteins with similar functions often share short 'motif' patterns — e.g., all kinases have a GXGXXG motif for ATP binding. We want to train a neural network to spot these motifs automatically and classify proteins into families (Kinase, GPCR, Protease, etc.).")
add_para(doc2, "Why? When scientists sequence a new genome, they find thousands of unknown proteins. Classifying them into families is the first step to understanding what they do. Traditional methods use HMMER or BLAST; deep learning can be more sensitive for distant relatives.")

add_step(doc2, 1, "Generate Synthetic Protein Sequences",
    "We create fake protein sequences (200-300 amino acids long) for 6 families. Each family has 2-3 defining motifs embedded in random background sequence. For example, Kinase gets GXGXXG + VAVK + HRDL embedded at random positions.",
    "generate_family_sequence(motifs, seq_length)",
    "We need labeled data to train a supervised classifier. Synthetic data lets us control exactly which motifs characterize each family. Later, we can replace this with real UniProt data.")

add_step(doc2, 2, "One-Hot Encode the Sequences",
    "Each amino acid becomes a vector of 20 zeros with a single 1 at its position. 'A' becomes [1,0,0,...,0], 'C' becomes [0,1,0,...,0], etc. A protein of length 250 becomes a 250x20 matrix. We pad or trim all proteins to the same length.",
    "one_hot_encode(seq, max_len=250)",
    "Neural networks need numerical input. One-hot encoding preserves the identity of each amino acid without implying any ordering (unlike numbers 1-20, which would suggest A < C < D, which is biologically meaningless).")

add_step(doc2, 3, "Build the 1D CNN",
    "The network has: Conv1D layer (64 filters, looking at 5-AA windows) → MaxPool → Conv1D (128 filters, 3-AA windows) → GlobalMaxPool → Dropout → Dense(64) → Dense(6 classes). Each filter learns to detect a specific motif pattern.",
    "class ProteinCNN(nn.Module)",
    "We use 1D CNN because: (1) Conv filters slide along the sequence detecting motifs anywhere (position-independent, just like real motifs), (2) Multiple stacked layers learn hierarchical patterns (simple motifs → motif combinations), (3) GlobalMaxPool picks the strongest motif hit per filter, making it robust to insertion/deletion shifts.")

add_step(doc2, 4, "Train the Network",
    "We feed batches of 32 proteins through the network, calculate how wrong the predictions are (CrossEntropyLoss), and adjust the weights using Adam optimizer. We repeat for 20 epochs (full passes through the data).",
    "nn.CrossEntropyLoss(), optim.Adam(model.parameters(), lr=1e-3)",
    "Training is where the network learns its motif detectors. Each epoch, the Conv filters adjust slightly to become better at spotting family-specific patterns. We track both training and test accuracy to watch for overfitting.")

add_step(doc2, 5, "Evaluate & Interpret",
    "We compute final test accuracy, per-family precision/recall/F1, and a confusion matrix. We also plot training curves (accuracy vs epoch) to see how learning progressed.",
    "accuracy_score(), classification_report(), confusion_matrix()",
    "Per-family metrics tell us which families are easy to classify (high F1) and which get confused (e.g., Kinase vs Helicase, since both bind ATP). The confusion matrix reveals these patterns at a glance.")

add_h(doc2, "What To Expect")
add_para(doc2, "The model should reach >95% accuracy since each family has distinct synthetic motifs. Training converges within 10-15 epochs. The confusion matrix may show confusion between families that share similar motifs (e.g., ATP-binding families). Output figure: project2_cnn_results.png.")

doc2.save(os.path.join(OUTPUT_DIR, 'Project2_Protein_CNN_Explained.docx'))
print("2/4 done")


# ════════════════════════════════════════════════════════════════════
# PROJECT 3
# ════════════════════════════════════════════════════════════════════
doc3 = Document()
doc3.add_heading('Project 3: Mutation Impact Prediction with ESM-2', 0)
add_para(doc3, 'A step-by-step walkthrough of using Meta\'s protein language model to predict whether a mutation destabilizes a protein.')

add_h(doc3, "The Big Picture — What & Why")
add_para(doc3, "When a single amino acid in a protein changes (a mutation), the protein might fall apart (destabilize), stop working, or cause disease. We want to predict which mutations are dangerous without doing expensive lab experiments.")
add_para(doc3, "Why? This is crucial for: interpreting genetic test results (is this mutation in your genome harmful?), designing better proteins (which mutations make a therapeutic antibody more stable?), and tracking viral evolution (will this Spike mutation make SARS-CoV-2 more fit?).")

add_step(doc3, 1, "Load Pretrained ESM-2 Model",
    "We download Meta's ESM-2 model (tiny version: 6 layers, 8 million parameters) from HuggingFace. This model was trained on 65 million protein sequences — it has learned the 'grammar' of proteins: which amino acids tend to appear together.",
    "AutoModelForMaskedLM.from_pretrained('facebook/esm2_t6_8M_UR50D')",
    "We don't train from scratch — that would take weeks on supercomputers. Instead, we use a pretrained model that already 'understands' proteins. This is called transfer learning, and it's the standard approach in modern AI.")

add_step(doc3, 2, "Create Mutation Dataset",
    "We take a reference protein (lambda repressor, 100 amino acids) and generate 300 single-point mutations (changing one amino acid at a time). Each mutation gets a simulated stability score (DDG) based on how different the new amino acid is from the original.",
    "blosum62_score(), mutation_stability_score()",
    "We need mutations with known stability effects to validate our method. In real research, you'd use ProTherm database (20,000+ experimental DDG values) or ClinVar (pathogenic vs benign human variants).")

add_step(doc3, 3, "Zero-Shot Mutation Scoring",
    "For each mutation, we: (1) Mask the mutated position in the protein sequence, (2) Ask ESM-2 to predict the probability of the original amino acid, (3) Ask ESM-2 to predict the probability of the mutant amino acid, (4) Score = log P(mutant) - log P(wild-type). A negative score means 'this mutation looks unlikely to the model'.",
    "score_mutation_simple(sequence, position, mutant_aa, model, tokenizer)",
    "This is called 'zero-shot' because we don't train anything — we just ask the pretrained model. The logic: if ESM-2 has seen millions of similar protein sequences and the mutant amino acid is rarely seen at this position, the mutation is probably destabilizing. No training data needed!")

add_step(doc3, 4, "Evaluate Zero-Shot Predictions",
    "We compare the ESM-2 scores against our simulated DDG values. A good model gives negative scores for destabilizing mutations. We compute: ROC-AUC (how well does it separate stabilizing from destabilizing?), Pearson correlation (does the score magnitude match DDG magnitude?).",
    "roc_auc_score(), mean_squared_error(), pearsonr()",
    "This tells us how much evolutionary information ESM-2 captured. Even without any training, these models typically achieve ROC-AUC ~0.7-0.8 for mutation effect prediction — comparable to physics-based methods like FoldX!")

add_step(doc3, 5, "Fine-Tune a Regression Head",
    "We extract per-position embeddings from ESM-2 (a 320-number vector for each amino acid). We freeze the ESM-2 weights and train a small neural network (320 → 64 → 32 → 1) to predict DDG from these embeddings.",
    "get_embeddings(), class RegressorHead(nn.Module)",
    "Fine-tuning adapts the general protein knowledge to our specific task. The frozen ESM-2 acts as a feature extractor, and the small regression head learns the mapping from 'general protein features' to 'stability prediction'. This usually improves correlation from r~0.3 to r~0.6.")

add_h(doc3, "What To Expect")
add_para(doc3, "Zero-shot: ROC-AUC ~0.65-0.85, Pearson r ~0.2-0.4 with DDG. Fine-tuned: r improves to ~0.5-0.7. The key insight: ESM-2 'knows' protein stability without ever being trained on it, because evolution has already done the experiment. Output figure: project3_mutation_esm2_results.png.")

doc3.save(os.path.join(OUTPUT_DIR, 'Project3_Mutation_ESM2_Explained.docx'))
print("3/4 done")


# ════════════════════════════════════════════════════════════════════
# PROJECT 4
# ════════════════════════════════════════════════════════════════════
doc4 = Document()
doc4.add_heading('Project 4: Viral Host Taxonomy Prediction with SVM', 0)
add_para(doc4, 'A step-by-step walkthrough of predicting which host a virus infects based on its genome sequence — based on Young et al. (2020).')

add_h(doc4, "The Big Picture — What & Why")
add_para(doc4, "When scientists discover a new virus (e.g., from a bat swab or wastewater sample), they don't always know what host it infects. But the virus's genome carries signatures of its host — because the host's immune system and cellular environment shape how the virus evolves.")
add_para(doc4, "Why? This matters for: pandemic preparedness (does this new virus have pandemic potential?), identifying zoonotic spillover events, understanding viral host range, and characterizing viruses from metagenomic sequencing where the host is unknown.")
add_para(doc4, "This project comes from Young et al. (2020) PLoS Computational Biology, featured in Alam et al. (2025) BMC Biology as a key SVM case study.")

add_step(doc4, 1, "Generate Viral Genomes with Host Signatures",
    "We create 800 synthetic viral genomes (5000 nt each, 200 per host) with different compositional biases: Mammalian viruses suppress CpG (immune evasion via ZAP protein), Plant viruses have higher GC, Bacteriophages are AT-rich. Each genome gets host-specific signature motifs inserted.",
    "generate_viral_genome(length, host_class)",
    "Each host type has a different biological environment that shapes viral evolution. Mammals have ZAP protein that targets CpG — so mammalian viruses evolve to avoid CpG. Plants don't have ZAP — their viruses don't suppress CpG. These differences become detectable patterns in the viral genome.")

add_step(doc4, 2, "Extract 340 Features from Each Genome",
    "We compute multiple feature types: (a) k=3 k-mer frequencies (64 numbers — trinucleotide composition), (b) k=4 k-mer frequencies (256 numbers — tetranucleotide composition), (c) Dinucleotide observed/expected ratios (16 numbers — which pairs appear more/less than expected), (d) GC content + skew (3 numbers), (e) Transition/Transversion ratio (1 number).",
    "extract_kmer_freqs(), extract_dinucleotide_bias(), extract_gc_features(), extract_titv_ratio()",
    "ML models need numerical vectors, not raw genomes. k-mers capture sequence composition at different scales. Dinucleotide bias captures CpG suppression (our key biological signal). GC features capture base composition. Ti/Tv ratio captures mutation bias. Each feature type captures a different biological signal.")

add_step(doc4, 3, "Standardize Features & Train SVM",
    "We standardize all features (subtract mean, divide by std dev) so they're on the same scale. Then we train an SVM with linear kernel (same method as Young et al. 2020). The SVM finds the best hyperplane separating the 4 host classes in the 340-dimensional feature space.",
    "StandardScaler(), SVC(kernel='linear', C=1.0)",
    "SVM is chosen because: (1) It works well with high-dimensional data (340 features, 800 samples), (2) The linear kernel makes it interpretable (we can see which features matter), (3) It's the method used in the original paper. Feature standardization is critical — SVM assumes all features have similar ranges.")

add_step(doc4, 4, "Train XGBoost for Comparison",
    "We also train an XGBoost classifier (200 trees, max_depth=6). XGBoost is a gradient-boosted tree ensemble that often outperforms SVM on tabular data. This lets us compare two different ML approaches on the same biological problem.",
    "XGBClassifier(n_estimators=200, max_depth=6)",
    "Comparing models is good scientific practice. SVM finds linear boundaries in the feature space; XGBoost builds trees that capture non-linear interactions. If both agree on which features matter, we can be more confident in the biology.")

add_step(doc4, 5, "Evaluate & Find Important Features",
    "We compute: Accuracy, per-class F1, confusion matrix (for both models), ROC-AUC (one-vs-rest for multi-class). We also check which features are most important — in SVM by coefficient magnitude, in XGBoost by built-in feature importance.",
    "accuracy_score(), classification_report(), roc_auc_score()",
    "Accuracy tells us overall performance, but per-class metrics reveal which hosts are easy/hard to predict. Feature importance is where the biology happens — if CpG-related features are top-ranked, it validates the ZAP hypothesis. If GC content dominates, the signal is different.")

add_step(doc4, 6, "Visualize Host Clustering",
    "We plot the genomes in GC-CpG space to see if they naturally cluster by host class. This is a biological validation step — if the data clusters by host in just 2 dimensions, the 340-feature ML model should easily separate them.",
    "plt.scatter(gc_content, cpg_per_100nt, c=host_color)",
    "This visualization confirms the biological premise is real. If mammalian viruses are in the low-CpG corner and plant viruses in the high-GC corner, we can literally SEE why the ML model works. This builds trust in the predictions.")

add_h(doc4, "What To Expect")
add_para(doc4, "Both SVM and XGBoost should achieve ~85-95% accuracy. CpG-related features (especially the 'CG' dinucleotide ratio) will be among the top predictors. The GC-CpG scatter plot will show clear host-specific clustering. Output: 4 figures in project4_output/.")
add_para(doc4, "Key biological insight: The model works because mammalian immune systems (ZAP protein) force viruses to avoid CpG. No ZAP in bacteria → no CpG suppression in phages. The virus genome is a 'fossil record' of its evolutionary history with its host.")

doc4.save(os.path.join(OUTPUT_DIR, 'Project4_Viral_Host_SVM_Explained.docx'))
print("4/4 done")
print(f"\nAll files saved to: {OUTPUT_DIR}")
